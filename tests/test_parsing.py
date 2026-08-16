"""Parser contract tests (BUILD_PLAN D4-M2, scenarios S1/S2)."""

from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

import pytest
from fpdf import FPDF
from PIL import Image

from ingestion.models import FormatKind
from ingestion.parsing import DocumentAIParser, LocalParser, Parser, UnsupportedFormatError

_DATA = Path(__file__).resolve().parent.parent / "data" / "acme_robotics"
_PINNED = datetime(2026, 7, 1, 9, 0, 0, tzinfo=UTC)


def _image_only_pdf_bytes() -> bytes:
    img = Image.new("RGB", (420, 240), "white")
    png = BytesIO()
    img.save(png, format="PNG")
    png.seek(0)
    pdf = FPDF()
    pdf.creation_date = _PINNED
    pdf.add_page()
    pdf.image(png)
    return bytes(pdf.output())


class TestLocalParser:
    def test_contract_pdf_text_and_metadata_fields(self) -> None:
        blob = (_DATA / "contract_customer_x.pdf").read_bytes()
        doc = LocalParser().parse(blob, "contract_customer_x.pdf", "deal-falcon")
        assert doc.text is not None
        assert "Change of Control" in doc.text
        assert doc.document_id == "contract_customer_x.pdf"
        assert doc.deal_id == "deal-falcon"
        assert doc.format.kind is FormatKind.NATIVE_PDF
        meta = doc.metadata
        assert meta["document_id"] == "contract_customer_x.pdf"
        assert meta["deal_id"] == "deal-falcon"
        assert meta["mime"] == "application/pdf"
        assert meta["needs_ocr"] is False
        assert meta["ingestion_timestamp"]

    def test_scanned_pdf_honest_text_none(self) -> None:
        doc = LocalParser().parse(_image_only_pdf_bytes(), "scanned_invoice.pdf", "deal-falcon")
        assert doc.text is None
        assert doc.format.needs_ocr is True
        assert doc.chunks == ()
        assert doc.metadata["needs_ocr"] is True

    def test_xlsx_sheets_parsed_to_tables(self) -> None:
        blob = (_DATA / "financials_fy27.xlsx").read_bytes()
        doc = LocalParser().parse(blob, "financials_fy27.xlsx", "deal-falcon")
        assert doc.text is not None
        assert "Meridian Logistics" in doc.text
        assert doc.tables, "xlsx must yield at least one table"
        first = doc.tables[0]
        assert "Customer" in first.headers
        flat = " | ".join(cell for row in first.rows for cell in row)
        assert "0.183" in flat
        assert first.name == "FY27 Projected Revenue"

    def test_docx_tmp_fixture_parsed(self) -> None:
        from docx import Document

        source = Document()
        source.add_paragraph("Fleet maintenance window moved to next quarter.")
        table = source.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "System"
        table.cell(0, 1).text = "Owner"
        table.cell(1, 0).text = "TitanBridge"
        table.cell(1, 1).text = "Ops"
        buf = BytesIO()
        source.save(buf)
        doc = LocalParser().parse(buf.getvalue(), "memo_fleet_operations.docx", "deal-falcon")
        assert doc.text is not None
        assert "Fleet maintenance window" in doc.text
        assert doc.tables and doc.tables[0].rows == (("TitanBridge", "Ops"),)

    def test_eml_stdlib_parsed(self) -> None:
        eml = (
            b"From: dana.whitfield@example.com\r\n"
            b"To: deal-team@example.com\r\n"
            b"Subject: Fleet maintenance window\r\n"
            b"\r\n"
            b"Maintenance moves to next quarter.\r\n"
        )
        doc = LocalParser().parse(eml, "memo.eml", "deal-falcon")
        assert doc.text is not None
        assert "Subject: Fleet maintenance window" in doc.text
        assert "Maintenance moves to next quarter." in doc.text

    def test_unknown_format_raises(self) -> None:
        with pytest.raises(UnsupportedFormatError):
            LocalParser().parse(b"\x00\x01garbage", "mystery.bin", "deal-falcon")

    def test_documentai_parser_refuses_without_enable_flag(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.delenv("DILIGENCE_DOCAI_ENABLED", raising=False)
        parser = DocumentAIParser(project_id="diligence-room", processor_id="abc123")
        blob = (_DATA / "contract_customer_x.pdf").read_bytes()
        with pytest.raises(RuntimeError, match="DILIGENCE_DOCAI_ENABLED"):
            parser.parse(blob, "contract_customer_x.pdf", "deal-falcon")

    def test_documentai_parser_satisfies_parser_protocol(self) -> None:
        parser = DocumentAIParser(project_id="diligence-room", processor_id="abc123")
        assert isinstance(parser, Parser)
        assert isinstance(LocalParser(), Parser)
