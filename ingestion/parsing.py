"""Document parsing (BUILD_PLAN D4-M2).

``LocalParser`` is the offline path (pypdf / openpyxl / python-docx / stdlib
email). ``DocumentAIParser`` is the live OCR path with the identical
``ParsedDoc`` contract, hard-guarded behind ``DILIGENCE_DOCAI_ENABLED=1`` so
offline tests can never touch the network. Scanned documents parse to
``text=None`` — text is never fabricated; routing the OCR path is the live
window's job.
"""

from __future__ import annotations

import os
from datetime import UTC, datetime
from email.parser import BytesParser
from email.policy import default as email_policy
from io import BytesIO
from typing import Protocol, runtime_checkable

from ingestion.formats import detect_format
from ingestion.models import FormatKind, ParsedDoc, TableData


class UnsupportedFormatError(ValueError):
    """Raised when a blob matches no parseable format class."""


@runtime_checkable
class Parser(Protocol):
    def parse(self, blob: bytes, document_id: str, deal_id: str) -> ParsedDoc: ...


def _cell_str(value: object) -> str:
    return "" if value is None else str(value)


def _pdf_text(blob: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(blob))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n".join(page for page in pages if page)


def _xlsx_parse(blob: bytes) -> tuple[str, tuple[TableData, ...]]:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(blob), data_only=True)
    text_lines: list[str] = []
    tables: list[TableData] = []
    for sheet in workbook.worksheets:
        rows = [
            tuple(_cell_str(cell) for cell in row)
            for row in sheet.iter_rows(values_only=True)
            if not all(cell is None for cell in row)
        ]
        text_lines.extend("\t".join(row) for row in rows)
        if rows:
            tables.append(TableData(name=sheet.title, headers=rows[0], rows=tuple(rows[1:])))
    return "\n".join(text_lines), tuple(tables)


def _dedup_row_cells(row: object) -> tuple[str, ...]:
    """Row texts, deduplicated — python-docx repeats merged cells."""
    seen: set[int] = set()
    cells: list[str] = []
    for cell in row.cells:  # type: ignore[attr-defined]
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        cells.append(cell.text.strip())
    return tuple(cells)


def _docx_parse(blob: bytes) -> tuple[str, tuple[TableData, ...]]:
    from docx import Document

    document = Document(BytesIO(blob))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables: list[TableData] = []
    for index, table in enumerate(document.tables):
        rows = tuple(_dedup_row_cells(row) for row in table.rows)
        if rows:
            tables.append(TableData(name=f"table:{index}", headers=rows[0], rows=tuple(rows[1:])))
    return "\n".join(paragraphs), tuple(tables)


def _eml_text(blob: bytes) -> str:
    message = BytesParser(policy=email_policy).parse(BytesIO(blob))
    header_lines = [
        f"{name}: {message[name]}" for name in ("From", "To", "Subject", "Date") if message[name]
    ]
    body_part = message.get_body(preferencelist=("plain",))
    body = body_part.get_content().strip() if body_part is not None else ""
    return "\n".join([*header_lines, "", body])


class LocalParser:
    """Offline parser: honest, deterministic, no network."""

    def parse(self, blob: bytes, document_id: str, deal_id: str) -> ParsedDoc:
        info = detect_format(blob, document_id)
        text: str | None
        tables: tuple[TableData, ...] = ()
        if info.kind is FormatKind.NATIVE_PDF:
            text = _pdf_text(blob)
        elif info.kind is FormatKind.XLSX:
            text, tables = _xlsx_parse(blob)
        elif info.kind is FormatKind.DOCX:
            text, tables = _docx_parse(blob)
        elif info.kind is FormatKind.EML:
            text = _eml_text(blob)
        elif info.kind in (FormatKind.SCANNED_PDF, FormatKind.IMAGE):
            text = None
        else:
            raise UnsupportedFormatError(f"cannot parse {document_id!r}: {info.reason}")
        metadata: dict[str, object] = {
            "document_id": document_id,
            "deal_id": deal_id,
            "format": info.kind.value,
            "mime": info.mime,
            "needs_ocr": info.needs_ocr,
            "ingestion_timestamp": datetime.now(UTC).isoformat(),
            "source": "local-parser",
        }
        return ParsedDoc(
            document_id=document_id,
            deal_id=deal_id,
            format=info,
            text=text,
            tables=tables,
            chunks=(),
            metadata=metadata,
        )


_DOC_AI_FLAG = "DILIGENCE_DOCAI_ENABLED"


def _layout_text(layout: object, text: str) -> str:
    """Join ALL text_anchor segments (half-open indices) — see Day-4 research."""
    anchor = getattr(layout, "text_anchor", None)
    if anchor is None:
        return ""
    segments = getattr(anchor, "text_segments", None) or []
    return "".join(
        text[int(getattr(seg, "start_index", 0) or 0) : int(getattr(seg, "end_index", 0) or 0)]
        for seg in segments
    )


class DocumentAIParser:
    """Live OCR/parser (Document AI), same ParsedDoc contract, flag-guarded.

    Region-pinned via ``location`` (endpoint ``{location}-documentai``);
    ``processor_type`` must be ``FORM_PARSER_PROCESSOR`` (or layout) for
    tables — the OCR processor returns none.
    """

    def __init__(
        self,
        project_id: str,
        processor_id: str,
        location: str = "us",
        processor_version: str = "pretrained-ocr-v2.0-2023-06-02",
        processor_type: str = "OCR_PROCESSOR",
    ) -> None:
        self.project_id = project_id
        self.processor_id = processor_id
        self.location = location
        self.processor_version = processor_version
        self.processor_type = processor_type

    def parse(self, blob: bytes, document_id: str, deal_id: str) -> ParsedDoc:
        if os.environ.get(_DOC_AI_FLAG) != "1":
            raise RuntimeError(
                "Document AI disabled offline; set DILIGENCE_DOCAI_ENABLED=1 (live window only)"
            )
        info = detect_format(blob, document_id)
        if info.kind is FormatKind.UNKNOWN:
            raise UnsupportedFormatError(f"cannot route {document_id!r}: {info.reason}")
        from google.api_core.client_options import ClientOptions
        from google.cloud import documentai

        client = documentai.DocumentProcessorServiceClient(
            client_options=ClientOptions(api_endpoint=f"{self.location}-documentai.googleapis.com")
        )
        name = client.processor_version_path(
            self.project_id, self.location, self.processor_id, self.processor_version
        )
        request = documentai.ProcessRequest(
            name=name,
            raw_document=documentai.RawDocument(content=blob, mime_type=info.mime),
        )
        result = client.process_document(request=request)
        document = result.document
        full_text: str = document.text or ""
        tables: list[TableData] = []
        for page_index, page in enumerate(document.pages):
            for table_index, table in enumerate(page.tables):
                header_rows = [
                    tuple(_layout_text(cell.layout, full_text).strip() for cell in row.cells)
                    for row in table.header_rows
                ]
                body_rows = tuple(
                    tuple(_layout_text(cell.layout, full_text).strip() for cell in row.cells)
                    for row in table.body_rows
                )
                tables.append(
                    TableData(
                        name=f"page:{page_index}:table:{table_index}",
                        headers=header_rows[0] if header_rows else (),
                        rows=body_rows,
                    )
                )
        metadata: dict[str, object] = {
            "document_id": document_id,
            "deal_id": deal_id,
            "format": info.kind.value,
            "mime": info.mime,
            "needs_ocr": False,
            "ingestion_timestamp": datetime.now(UTC).isoformat(),
            "source": f"documentai:{self.processor_type}",
        }
        return ParsedDoc(
            document_id=document_id,
            deal_id=deal_id,
            format=info,
            text=full_text or None,
            tables=tuple(tables),
            chunks=(),
            metadata=metadata,
        )
